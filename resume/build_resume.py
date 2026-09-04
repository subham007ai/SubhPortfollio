# -*- coding: utf-8 -*-
"""
ATS-friendly resume generator for Subham Sarangi.
Content sourced from D:/personal-4 portfolio (content/*.ts), then corrected by
Subham: languages he does not actually use removed, AI-assisted work stated plainly.

ATS rules honoured:
  - single column, no tables, no text boxes, no images, no icons
  - no content in Word headers/footers (parsers often drop them)
  - standard section headings (Summary / Skills / Experience / Projects / Education ...)
  - standard font (Calibri), plain bullet glyph, hanging indents
  - contact details as plain text in the document body
  - ASCII-safe currency ("INR 20,000" not the rupee sign)
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\personal-4\resume"
DOCX = os.path.join(OUT_DIR, "Subham_Sarangi_Resume.docx")

FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x50, 0x50, 0x50)

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
normal.font.color.rgb = INK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

for s in doc.sections:
    s.top_margin = Inches(0.5)
    s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.6)
    s.right_margin = Inches(0.6)


def para(space_before=0, space_after=0, align=None, left=0, hanging=None):
    p = doc.add_paragraph()
    f = p.paragraph_format
    f.space_before = Pt(space_before)
    f.space_after = Pt(space_after)
    f.line_spacing = 1.0
    if left:
        f.left_indent = Inches(left)
    if hanging:
        f.first_line_indent = Inches(-hanging)
    if align:
        p.alignment = align
    return p


def run(p, text, size=10, bold=False, italic=False, color=INK, caps=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    return r


def hyperlink(p, url, text=None, size=9.5):
    """Real hyperlink whose *display text* is the plain URL -> parses cleanly in ATS."""
    text = text or url
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rPr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "1A1A1A")
    rPr.append(col)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    link.append(new_run)
    p._p.append(link)


def rule(p):
    """Thin bottom border on a paragraph (safe for ATS - not a table/shape)."""
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "9A9A9A")
    borders.append(bottom)
    pPr.append(borders)


def heading(text, first=False):
    p = para(space_before=0 if first else 7, space_after=3)
    run(p, text.upper(), size=10.5, bold=True)
    rule(p)


def entry(title, meta_right=None, sub=None, sub_right=None):
    """Role/project line + optional right-aligned date, then an italic sub-line.

    keep_with_next stops a heading from being orphaned at the foot of a page.
    """
    p = para(space_before=4, space_after=0)
    p.paragraph_format.keep_with_next = True
    tab = p.paragraph_format.tab_stops
    tab.add_tab_stop(Inches(7.3), 2)  # 2 = RIGHT
    run(p, title, size=10.5, bold=True)
    if meta_right:
        run(p, "\t" + meta_right, size=9.5, color=GREY)
    if sub:
        p2 = para(space_before=0.5, space_after=1)
        p2.paragraph_format.keep_with_next = True
        t2 = p2.paragraph_format.tab_stops
        t2.add_tab_stop(Inches(7.3), 2)
        run(p2, sub, size=9.5, italic=True, color=GREY)
        if sub_right:
            run(p2, "\t" + sub_right, size=9.5, italic=True, color=GREY)


def bullet(text, bold_lead=None, keep=False):
    p = para(space_before=1, space_after=1, left=0.19, hanging=0.19)
    p.paragraph_format.keep_together = True
    if keep:
        p.paragraph_format.keep_with_next = True
    run(p, "\u2022  ", size=10)
    if bold_lead:
        run(p, bold_lead, size=10, bold=True)
    run(p, text, size=10)


def kv(label, value):
    p = para(space_before=1.5, space_after=1.5, left=0.95, hanging=0.95)
    run(p, label + ": ", size=10, bold=True)
    run(p, value, size=10)


# =====================================================================
# HEADER  (in body, not a Word header)
# =====================================================================
p = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "SUBHAM SARANGI", size=20, bold=True)

p = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "AI & Machine Learning Student  |  Builds Web and ML Projects Using AI Tools",
    size=10.5, color=GREY)

p = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Bhubaneswar, Odisha, India  |  +91 93376 82421  |  ", size=9.5)
hyperlink(p, "mailto:worksubhamsarangi@gmail.com", "worksubhamsarangi@gmail.com")

# =====================================================================
heading("Summary")
p = para(space_after=0)
run(
    p,
    "I'm a B.Tech Computer Science student at SOA University, Bhubaneswar, studying AI and Machine "
    "Learning (2024-2028). I build most of my projects using AI tools like Claude and ChatGPT - I "
    "write the prompts, read through what comes back, test it, fix what doesn't work, and keep going "
    "until the thing is finished and online. That's the part I've got genuinely good at. So far I've "
    "put six websites live, two of them paid client jobs, and worked on a face mask detection project "
    "where I handled the data side. My team came 2nd out of 400+ at SOA Ideathon 2025. I'm looking "
    "for an internship or entry-level role where I can keep building and learn the parts I don't "
    "know yet.",
    size=10,
)

# =====================================================================
heading("Skills")
kv("AI Tools", "Claude, ChatGPT and other AI tools; prompt writing and refining, using AI to "
               "build and debug real projects")
kv("Languages", "Python, HTML, CSS")
kv("Data & ML", "NumPy, Pandas, scikit-learn, TensorFlow/Keras, OpenCV, Matplotlib, Seaborn")
kv("Web", "Next.js, React, Tailwind CSS, Framer Motion")
kv("Backend & Infra", "Flask, FastAPI, Node.js, REST APIs - built with AI assistance")
kv("Tools", "Git, GitHub, VS Code, Jupyter, Vercel, Figma")

# =====================================================================
heading("Experience")

entry(
    "Web Development Intern",
    "2026",
    "Odisha International Short Film Festival (OISFF) - Bhubaneswar, India",
    "6-week internship",
)
bullet("Built the full website for Odisha's first international short film festival and got it "
       "finished inside a six-week internship. It's live at odishafilmsociety.in.")
bullet("Put together eight sections - mission, programs, film categories, gallery, news, a countdown "
       "timer and a film submission form - and made sure everything worked properly on phones, "
       "tablets and desktops.")
bullet("Worked out the whole look myself: Sanskrit-style serif headings, Konark chakra icons and a "
       "bronze and gold colour scheme, so it felt like a real cultural festival and not a template. "
       "Stack: Next.js, React, TypeScript, Tailwind CSS, Framer Motion.")

entry(
    "Freelance Web Developer",
    "2025 - Present",
    "Self-employed - Bhubaneswar, India",
)
bullet("Build and deliver bespoke web platforms for cultural institutions and creative ventures, "
       "handling projects from initial consultation through to putting the site online.")
bullet("Work out what a client actually needs from a non-technical brief, and explain my choices "
       "back to them in plain language.")

# =====================================================================
heading("Projects")

entry(
    "Face Mask Detection - Real-Time Detection System",
    "2025",
    "College group project (4 members) - Python, TensorFlow/Keras, OpenCV, Flask",
)
bullet("Handled the data side: checked how balanced the two classes were, built the image "
       "preprocessing pipeline, and set up the data generators used to train EfficientNetB0, "
       "MobileNetV2 and ResNet50.")
bullet("Compared all three models on a 1,726-image test set (863 images per class) using accuracy, "
       "precision, recall, F1 and ROC-AUC. MobileNetV2 came out at 99.59%.")
bullet("EfficientNetB0 scored a perfect 100%, which looked wrong, so we checked the train/test split "
       "for leakage before trusting it. We picked it for the final build mainly because it's much "
       "smaller - around 49 MB against ResNet50's 283 MB.")
bullet("The finished app runs in a browser through Flask and shows a live video feed with a box "
       "around the face and a confidence score, at roughly 30 FPS on a normal CPU with no graphics "
       "card needed.")

entry(
    "Sitora - Creative Agency Website",
    "2025",
    "Personal project - Next.js 15, React 19, TypeScript, Tailwind CSS 3, Framer Motion, Lenis, Zod",
)
bullet("Built a colour system where each section sets its own accent colour and the whole page shifts "
       "colour smoothly as you scroll down.")
bullet("Made a pricing calculator with three sliders and an INR/USD switch that passes the estimate "
       "straight through into the contact form.")
bullet("Added proper form validation and a hidden field that quietly blocks spam bots.")

entry(
    "Personal Portfolio Website",
    "2026",
    "Personal project - Next.js, React, TypeScript, Tailwind CSS, Framer Motion",
)
bullet("Built my own portfolio with dark and light modes, page transitions, an activity heatmap and "
       "a photo lightbox.")

# =====================================================================
heading("Achievements")
bullet("Came 2nd out of 400+ teams, working in a team of six. Prize: INR 20,000.",
       bold_lead="Runner-Up, SOA Ideathon 2025, SOA University (2025) - ")
bullet("Worked in a team of four on an idea for spotting illegal mining, built in 24 hours.",
       bold_lead="24-Hour Hackathon, NIT Foundation, Nalanda Institute of Technology (2026) - ")

# =====================================================================
heading("Education")
entry(
    "B.Tech, Computer Science and Engineering (Artificial Intelligence & Machine Learning)",
    "Oct 2024 - Jan 2028",
    "Siksha 'O' Anusandhan University, Bhubaneswar, Odisha",
)
entry(
    "Higher Secondary (Class XII), Science - 68%",
    "2024",
    "Bridgewell Global School, Bhubaneswar, Odisha",
)
entry(
    "Secondary (Class X) - 89%",
    "2022",
    "Bridgewell Global School, Bhubaneswar, Odisha",
)

# =====================================================================
heading("Certifications")
p = para(space_after=0)
run(p, "Python Full Course, GeeksforGeeks (Jul 2025)  |  Learn HTML & CSS, Scrimba (Jul 2025)  |  "
       "AI Tools Workshop, Be10x (Jun 2025)  |  Excel for Beginners, Great Learning (Oct 2024)",
    size=10)

# =====================================================================
heading("Languages")
p = para(space_after=0)
run(p, "English (Professional Working)  |  Hindi (Professional Working)  |  Odia (Native)", size=10)

os.makedirs(OUT_DIR, exist_ok=True)
doc.save(DOCX)
print("Saved:", DOCX)
